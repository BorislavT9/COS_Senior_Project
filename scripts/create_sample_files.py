"""
Create minimal sample PDF, DOCX, XLSX in sample_docs/ for spikes. Run from project root.
"""
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))


def create_sample_docx() -> None:
    from docx import Document
    doc = Document()
    doc.add_paragraph("Sample DOCX for Week 3 spike.")
    doc.add_paragraph("Invoice #INV-2024-001")
    out = ROOT / "sample_docs" / "sample.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print("Created", out)


def create_sample_xlsx() -> None:
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "Name"
    ws["B1"] = "Value"
    ws["A2"] = "Invoice"
    ws["B2"] = "INV-2024-001"
    out = ROOT / "sample_docs" / "sample.xlsx"
    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out))
    print("Created", out)


def create_sample_pdf() -> None:
    from pypdf import PdfWriter
    writer = PdfWriter()
    # Minimal PDF (blank page); pypdf has no high-level text API for writing
    writer.add_blank_page(width=72 * 4, height=72 * 4)
    writer.add_metadata({"/Title": "Sample PDF"})
    out = ROOT / "sample_docs" / "sample.pdf"
    out.parent.mkdir(parents=True, exist_ok=True)
    with open(out, "wb") as f:
        writer.write(f)
    print("Created", out)


if __name__ == "__main__":
    create_sample_docx()
    create_sample_xlsx()
    create_sample_pdf()
